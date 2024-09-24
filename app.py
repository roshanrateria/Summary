from flask import Flask, render_template, request, redirect, url_for, send_file, flash
import os
import pandas as pd
from scholarly import scholarly
import bibtexparser
from io import BytesIO
from docx import Document
import plotly
import plotly.express as px
import json
import functools
import datetime
from scholarly import ProxyGenerator

# Set up a ProxyGenerator object to use free proxies
# This needs to be done only once per session

app = Flask(__name__)
app.secret_key = 'your_secret_key'  # Replace with a secure key
app.config['UPLOAD_FOLDER'] = 'uploads'
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Global variable to store publication data
publication_data = pd.DataFrame()

def parse_bibtex(file_path):
    with open(file_path, encoding='utf-8') as bibtex_file:
        bib_database = bibtexparser.load(bibtex_file)
    records = []
    for entry in bib_database.entries:
        records.append({
            'title': entry.get('title', ''),
            'author': entry.get('author', ''),
            'year': entry.get('year', ''),
            'venue': entry.get('journal', entry.get('booktitle', '')),
            'type': 'Journal' if 'journal' in entry else 'Conference'
        })
    return pd.DataFrame(records)
@functools.cache
def fetch_scholarly_data(faculty_name):
    search_query = scholarly.search_author(faculty_name)
    try:
        author = next(search_query)
        author = scholarly.fill(author, sections=['publications'])
        publications = author.get('publications', [])
        records = []
        for pub in publications:
            pub_filled = scholarly.fill(pub)
            records.append({
                'title': pub_filled.get('bib', {}).get('title', ''),
                'author': ', '.join(pub_filled.get('bib', {}).get('author', [])),
                'year': pub_filled.get('bib', {}).get('pub_year', ''),
                'venue': pub_filled.get('bib', {}).get('journal', pub_filled.get('bib', {}).get('booktitle', '')),
                'type': 'Journal' if 'journal' in pub_filled.get('bib', {}) else 'Conference'
            })
        return pd.DataFrame(records)
    except StopIteration:
        return pd.DataFrame()

@app.route('/')
def home():
    return redirect(url_for('upload'))

@app.route('/upload', methods=['GET', 'POST'])
def upload():
    global publication_data
    if request.method == 'POST':
        if 'file' not in request.files:
            flash('No file part')
            return redirect(request.url)
        file = request.files['file']
        if file.filename == '':
            flash('No selected file')
            return redirect(request.url)
        if file:
            filename = f"{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}_{file.filename}"
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)
            # Determine file type
            if filename.endswith('.xlsx') or filename.endswith('.xls'):
                df = pd.read_excel(file_path)
            elif filename.endswith('.bib'):
                df = parse_bibtex(file_path)
            else:
                flash('Unsupported file format')
                return redirect(request.url)
            # Fetch additional data from scholarly
            faculty_names = df['Faculty Name'].unique()
            all_records = []
            for name in faculty_names:
                fetched_data = fetch_scholarly_data(name)
                print(fetched_data)
                fetched_data['Faculty Name'] = name
                all_records.append(fetched_data)
            if all_records:
                publication_data = pd.concat(all_records, ignore_index=True)
            else:
                publication_data = df
            flash('File successfully uploaded and processed')
            return redirect(url_for('report'))
    return render_template('upload.html')
import numpy as np
@app.route('/report')
def report():
    if publication_data.empty:
        flash('No data available. Please upload a file first.')
        return redirect(url_for('upload'))
    # Generate plots
    # Publications per Year
    publication_data['year']= publication_data['year'].replace('', np.nan)

    pub_year = publication_data['year'].dropna().astype(int)
    fig1 = px.histogram(pub_year, x=pub_year, nbins=20, title='Publications per Year')
    graph1 = json.dumps(fig1, cls=plotly.utils.PlotlyJSONEncoder)

    # Publications by Type
    pub_type = publication_data['type'].value_counts().reset_index()
    pub_type.columns = ['Type', 'Count']
    fig2 = px.pie(pub_type, names='Type', values='Count', title='Publications by Type')
    graph2 = json.dumps(fig2, cls=plotly.utils.PlotlyJSONEncoder)

    # Publications by Faculty
    pub_faculty = publication_data['Faculty Name'].value_counts().reset_index()
    pub_faculty.columns = ['Faculty Name', 'Count']
    fig3 = px.bar(pub_faculty, x='Faculty Name', y='Count', title='Publications by Faculty')
    graph3 = json.dumps(fig3, cls=plotly.utils.PlotlyJSONEncoder)

    return render_template('report.html', graph1=graph1, graph2=graph2, graph3=graph3,fac=publication_data['Faculty Name'].unique() )

@app.route('/export', methods=['GET'])
def export():
    if publication_data.empty:
        flash('No data to export')
        return redirect(url_for('upload'))
    export_format = request.args.get('format', 'excel')
    buffer = BytesIO()
    if export_format == 'excel':
        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
            publication_data.to_excel(writer, index=False, sheet_name='Publications')
        buffer.seek(0)
        return send_file(buffer, as_attachment=True, download_name='publication_records.xlsx', mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    elif export_format == 'word':
        doc = Document()
        doc.add_heading('Publication Records', 0)
        table = doc.add_table(rows=1, cols=len(publication_data.columns))
        hdr_cells = table.rows[0].cells
        for i, column in enumerate(publication_data.columns):
            hdr_cells[i].text = column
        for _, row in publication_data.iterrows():
            row_cells = table.add_row().cells
            for i, item in enumerate(row):
                row_cells[i].text = str(item)
        doc.save(buffer)
        buffer.seek(0)
        return send_file(buffer, as_attachment=True, download_name='publication_records.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    else:
        flash('Unsupported export format')
        return redirect(url_for('report'))

@app.route('/export/name', methods=['GET'])
def export_name():
    if publication_data.empty:
        flash('No data to export')
        return redirect(url_for('upload'))
    faculty_name = request.args.get('faculty')
    if not faculty_name:
        flash('No faculty name provided')
        return redirect(url_for('report'))
    filtered_data = publication_data[publication_data['Faculty Name'] == faculty_name]
    if filtered_data.empty:
        flash('No data found for the specified faculty')
        return redirect(url_for('report'))
    export_format = request.args.get('format', 'excel')
    buffer = BytesIO()
    if export_format == 'excel':
        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
            filtered_data.to_excel(writer, index=False, sheet_name='Publications')
        buffer.seek(0)
        filename = f'publication_records_{faculty_name}.xlsx'
        return send_file(buffer, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    elif export_format == 'word':
        doc = Document()
        doc.add_heading(f'Publication Records for {faculty_name}', 0)
        table = doc.add_table(rows=1, cols=len(filtered_data.columns))
        hdr_cells = table.rows[0].cells
        for i, column in enumerate(filtered_data.columns):
            hdr_cells[i].text = column
        for _, row in filtered_data.iterrows():
            row_cells = table.add_row().cells
            for i, item in enumerate(row):
                row_cells[i].text = str(item)
        doc.save(buffer)
        buffer.seek(0)
        filename = f'publication_records_{faculty_name}.docx'
        return send_file(buffer, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    else:
        flash('Unsupported export format')
        return redirect(url_for('report'))

@app.route('/export/year', methods=['GET'])
def export_year():
    publication_data['author'] = publication_data['author'].str.replace(',', '')

    if publication_data.empty:
        flash('No data to export')
        return redirect(url_for('upload'))
    start_year = request.args.get('start_year')
    end_year = request.args.get('end_year')
    if not start_year or not end_year:
        flash('Start year and end year must be provided')
        return redirect(url_for('report'))
    try:
        start_year = float(start_year)
        end_year = float(end_year)
    except ValueError:
        flash('Invalid year format')
        return redirect(url_for('report'))
    publication_data['year']= publication_data['year'].replace('', np.nan)
    publication_data['year'].dropna(inplace=True)
    print(publication_data['year'])
    try:
        filtered_data = publication_data[(publication_data['year'] >= start_year) & (publication_data['year'] <= end_year)]
        if filtered_data.empty:
            flash('No data found for the specified duration')
            return redirect(url_for('report'))
    except:
        flash('No data found for the specified duration')
        return redirect(url_for('report'))
    export_format = request.args.get('format', 'excel')
    buffer = BytesIO()
    if export_format == 'excel':
        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
            filtered_data.to_excel(writer, index=False, sheet_name='Publications')
        buffer.seek(0)
        filename = f'publication_records_{start_year}_{end_year}.xlsx'
        return send_file(buffer, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    elif export_format == 'word':
        doc = Document()
        doc.add_heading(f'Publication Records from {start_year} to {end_year}', 0)
        table = doc.add_table(rows=1, cols=len(filtered_data.columns))
        hdr_cells = table.rows[0].cells
        for i, column in enumerate(filtered_data.columns):
            hdr_cells[i].text = column
        for _, row in filtered_data.iterrows():
            row_cells = table.add_row().cells
            for i, item in enumerate(row):
                row_cells[i].text = str(item)
        doc.save(buffer)
        buffer.seek(0)
        filename = f'publication_records_{start_year}_{end_year}.docx'
        return send_file(buffer, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    else:
        flash('Unsupported export format')
        return redirect(url_for('report'))

if __name__ == '__main__':
    app.run(debug=True)
